import os
import uuid
import json
from docx import Document
from docx.shared import Inches
from random import shuffle
from models import db
from models.question import Question
from models.quiz import Quiz
from api.v1.views import app_views, needed, question_options
from flask import Flask, request, abort, jsonify
from utils import generate_code
from werkzeug.utils import secure_filename

UPLOAD_FOLDER = 'app_images/question_images/'
DOCUMENT_FOLDER = 'document/files/'

@app_views.route('/all-questions/<teacher_id>', methods=['GET'])
@app_views.route('/all-questions', methods=['GET'])
def get_allquestions(teacher_id=''):
    """gets all questions"""

    if teacher_id:
        questions = db.session.execute(db.select(Question).filter_by(teacher_id=teacher_id)).all()
    else:
        questions = db.session.execute(db.select(Question)).all()
    print(questions)
    if questions:
        new_questions = []
        for objs in questions:
            obj= objs[0]
            new_dict = {}
            options = []
            for key, value in obj.__dict__.items():
                if key in needed:
                    new_dict[key] = value
                if key in question_options:
                    options.append(value)
            shuffle(options)
            new_dict['options'] = options
            new_questions.append(new_dict)
        return jsonify(new_questions)
    else:
        return jsonify([])
    
@app_views.route('/create-new/<id>', methods=['POST'])
def create_new(id):
    """post questions from a teacher and
    creates a new quiz"""

    question_metadata = request.form
    files = request.files
    if not question_metadata:
        abort(404)

    print(question_metadata)
    duration = question_metadata.get('duration')
    question_list = question_metadata.get('questions')
    subject = question_metadata.get('Subject')
  
    question_list = json.loads(question_list)

    # creating the quiz object/model
    quiz = Quiz(id = uuid.uuid4())
    quiz.teacher_id = id
    quiz.duration = duration
    quiz.code = generate_code()
    quiz.subject = subject

    # creates a word document
    doc = Document()

    for n, question_dic in enumerate(question_list):
        question_dic['id'] = uuid.uuid4()
        question_dic['teacher_id'] = id
        question_dic['subject'] = subject

        # handles the saving of the images for the questions
        del question_dic['image']
        if 'image' + str(n) in files.keys():
            file = files.get('image' + str(n))
            if file:
                filename = secure_filename(file.filename)
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)
                question_dic['image'] = filepath

        question = Question(**question_dic)
        db.session.add(question)
        quiz.questions.append(question)
        db.session.commit()

        p = doc.add_paragraph('')
        p.add_run('Question {}'.format(n + 1)).bold = True
        doc.add_paragraph(question_dic['body'])
        if question_dic.get('image'):
            doc.add_picture(question_dic['image'], width=Inches(1.25), height=Inches(1.25))
        options = [question_dic['right_answer'],
                   question_dic['wrong_answer1'],
                   question_dic['wrong_answer2'],
                   question_dic['wrong_answer3'],
                   question_dic['wrong_answer4']]
        shuffle(options)
        doc.add_paragraph('A {}    B {}    C {}    D {}   E {}'.format(options[0],
                                                            options[1],
                                                            options[2],
                                                            options[3],
                                                            options[4]))
    docpath = DOCUMENT_FOLDER+'{}.docx'.format(quiz.id)
    docpdf = DOCUMENT_FOLDER+'{}.pdf'.format(quiz.id)
    doc.save(docpath)
    doc.save(docpdf)
    quiz.doc = docpath
    quiz.pdf = docpdf
    db.session.add(quiz)
    db.session.commit()
    
    result = {'quiz_id': quiz.id, 'message': 'quiz created', 'code': quiz.code, 'docFile': quiz.doc}
    return jsonify(result)

@app_views.route('/create-existing/<id>', methods=['POST'])
def create_existing(id):
    """creates a quiz from pre existing questions"""

    # gets a metadata from the get json, it contains the questions and the duration for the quiz
    # it is a dictionary
    question_metadata = request.get_json()
    print(question_metadata)
    if not question_metadata:
        abort(404)
    question_id_list = question_metadata.get('ids')
    duration = question_metadata.get('duration')
    subject = question_metadata.get('Subject')
    if not question_id_list:
        abort(404)
    quiz = Quiz(id = uuid.uuid4())
    quiz.teacher_id = id
    quiz.duration = duration
    quiz.code = generate_code()
    quiz.subject = subject
    

    doc = Document()

    for n, q_id in enumerate(question_id_list):
        question = db.get_or_404(Question, q_id)
        quiz.questions.append(question)
        db.session.commit()

        p = doc.add_paragraph('')
        p.add_run('Question {}'.format(n + 1)).bold = True
        doc.add_paragraph(question.body)
        if question.image:
            doc.add_picture(question.image, width=Inches(2.0), height=Inches(1.25))
        options = [question.right_answer,
                   question.wrong_answer1,
                   question.wrong_answer2,
                   question.wrong_answer3,
                   question.wrong_answer4]
        shuffle(options)
        doc.add_paragraph('A.{}    B.{}    C.{}    D.{}    E.{}'.format(options[0],
                                                            options[1],
                                                            options[2],
                                                            options[3],
                                                            options[4]))
    docpath = DOCUMENT_FOLDER+'{}.docx'.format(quiz.id)
    docpdf = DOCUMENT_FOLDER+'{}.pdf'.format(quiz.id)
    doc.save(docpath)
    doc.save(docpdf)
    quiz.doc = docpath
    quiz.pdf = docpdf
    db.session.add(quiz)
    db.session.commit()
    
    result = {'quiz_id': quiz.id, 'message': 'quiz created', 'code': quiz.code, 'docFile': quiz.doc}
    return jsonify(result)
    


    